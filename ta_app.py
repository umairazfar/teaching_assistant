import streamlit as st
import pandas as pd
import json
import os
import io
import zipfile
from google import genai
import google.api_core.exceptions as google_exceptions
from google.genai import types
from fpdf import FPDF
import docx
from pypdf import PdfReader
import logging
import time
from dotenv import load_dotenv
from st_supabase_connection import SupabaseConnection
from supabase import create_client

load_dotenv()


# 1. Setup Connection
supabase = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

def login_screen():
    st.title("🍎 AI Teaching Assistant Login")
    
    tab1, tab2 = st.tabs(["Magic Link", "Google Login"])
    
    with tab1:
        email = st.text_input("Enter your university email")
        if st.button("Send Magic Link"):
            res = supabase.auth.sign_in_with_otp({"email": email})
            st.success("Check your email for the login link!")

    with tab2:
        if st.button("Sign in with Google"):
            # This redirects the user to Google
            res = supabase.auth.sign_in_with_oauth({
                "provider": "google",
                "options": { "redirect_to": "http://localhost:8501" }
            })
            # Streamlit logic to handle the redirect follows here...

# ==========================================
# 1. DATABASE & AUTHENTICATION HELPERS
# ==========================================
# This helper handles the Supabase Connection
conn = st.connection("supabase", type=SupabaseConnection)

def get_user_credits(email):
    """Fetches credits from the profile_credits table we linked to Auth."""
    res = conn.table("profile_credits").select("credits").eq("email", email).execute()
    if len(res.data) == 0:
        # Trigger in Supabase usually handles this, but as a fallback:
        return 0
    return res.data[0]["credits"]

def deduct_credit(email, amount=1):
    """Subtracts credits after successful grading."""
    current = get_user_credits(email)
    if current >= amount:
        new_total = current - amount
        conn.table("profile_credits").update({"credits": new_total}).eq("email", email).execute()
        return True
    return False

# --- DYNAMIC AUTH LOGIC (No more dev_mode flag) ---
# Check if Supabase has an active session for this browser
user_info = st.context.user  # Streamlit's built-in way to access logged-in user

if not user_info:
    st.set_page_config(page_title="AI Teaching Assistant", page_icon="🎓")
    st.title("👨‍🏫 AI Teaching Assistant")
    st.info("Please log in with Google to access the grading engine.")
    # This button triggers the Supabase Google OAuth handshake
    st.button("Log in with Google", on_click=st.login)
    st.stop()

# If we reach here, the user is logged in
user_email = user_info.email
user_name = user_info.get("name", "Professor")
balance = get_user_credits(user_email)

# ==========================================
# 2. UI CONFIG & STYLING
# ==========================================
st.set_page_config(page_title="AI Teaching Assistant", layout="wide", page_icon="🎓")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
    [data-testid="stColumn"] {
        background-color: rgba(128, 128, 128, 0.05);
        padding: 2rem;
        border-radius: 12px;
    }
    div.stButton > button:first-child {
        background-color: #1A73E8;
        color: white;
        font-weight: 600;
    }
    </style>
    """, unsafe_allow_html=True)

# Initialize Session State
if "grading_data" not in st.session_state: st.session_state.grading_data = None
if "raw_json_list" not in st.session_state: st.session_state.raw_json_list = []
if "zip_bytes" not in st.session_state: st.session_state.zip_bytes = None
if "confirm_no_sol" not in st.session_state: st.session_state.confirm_no_sol = False
if "run_evaluation" not in st.session_state: st.session_state.run_evaluation = False
if "file_states" not in st.session_state: st.session_state.file_states = {"rubric": None, "sol": None, "students": 0}
if "report_index" not in st.session_state: st.session_state.report_index = 0
if "show_confirm_dialog" not in st.session_state: st.session_state.show_confirm_dialog = False

# ==========================================
# 3. CORE PROCESSING FUNCTIONS
# ==========================================
def extract_text(uploaded_file):
    fname = uploaded_file.name.lower()
    try:
        if fname.endswith('.zip'):
            combined_text = ""
            with zipfile.ZipFile(uploaded_file) as z:
                for file_info in z.infolist():
                    if not file_info.is_dir():
                        inner_fname = file_info.filename.lower()
                        with z.open(file_info) as f:
                            if inner_fname.endswith('.pdf'):
                                reader = PdfReader(f)
                                content = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
                            elif inner_fname.endswith('.docx'):
                                doc = docx.Document(f)
                                content = "\n".join([p.text for p in doc.paragraphs])
                            else:
                                content = f.read().decode("utf-8", errors="ignore")
                            combined_text += f"\n--- FILE: {file_info.filename} ---\n{content}\n"
            return combined_text
        elif fname.endswith(('.csv', '.xlsx')):
            df = pd.read_csv(uploaded_file) if fname.endswith('.csv') else pd.read_excel(uploaded_file)
            return df.to_string()
        elif fname.endswith('.pdf'):
            reader = PdfReader(uploaded_file)
            return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif fname.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return uploaded_file.read().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"Error reading {fname}: {str(e)}"

def create_pdf_report(student_data):
    pdf = FPDF()
    pdf.add_page()
    def clean_text(text):
        if not text: return ""
        replacements = {'\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"'}
        for char, rep in replacements.items(): text = str(text).replace(char, rep)
        return text.encode('latin-1', 'replace').decode('latin-1')

    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 15, f"Report: {clean_text(student_data.get('username'))}", ln=True, align='C')
    pdf.line(10, 25, 200, 25)
    pdf.ln(10)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, f"Score: {student_data.get('total_score', 0)} points", ln=True)
    pdf.ln(5)
    pdf.set_font("Arial", "", 11)
    pdf.multi_cell(0, 6, clean_text(student_data.get('overall_feedback', '')))
    for q in student_data.get("questions", []):
        pdf.ln(5)
        pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 8, f"Q{q.get('q_num')} | Score: {q.get('score')}", ln=True)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 5, clean_text(q.get('feedback', '')))
    return pdf.output(dest='S').encode('latin-1')

# ==========================================
# 4. SIDEBAR & MAIN UI
# ==========================================

# Define it once using the correct ID from your Stripe Dashboard
#stripe_payment_url = f"https://buy.stripe.com/test_3cI7sKaAU0VvgjA1UgaR200?prefilled_email={user_email}"

stripe_url = st.secrets.get("STRIPE_PAYMENT_URL", "https://buy.stripe.com/test_3cI7sKaAU0VvgjA1UgaR200")
stripe_payment_url = f"{stripe_url}?prefilled_email={user_email}"

with st.sidebar:
    st.header("🪙 Credits")
    st.metric("Balance", f"🪙 {balance}")
    # Use the consistent variable here
    st.link_button("Buy More Credits", stripe_payment_url, width='content')
    st.divider()
    
    # 2. Hardcoded API Key
    api_key = st.secrets["GEMINI_API_KEY"]

    model_choice = "gemini-2.5-flash" # or "gemini-1.5-flash" for the stable version

st.title("👨‍🏫 AI Teaching Assistant")
st.write(f"Logged in as: **{user_email}**")

col1, col2 = st.columns(2, gap="large")
with col1:
    st.subheader("📁 Reference Materials")
    rubric_file = st.file_uploader("Upload Rubric", key="rubric")
    sol_file = st.file_uploader("Upload Solution", type=['py', 'zip', 'pdf', 'docx', 'txt'], key="sol")

with col2:
    st.subheader("📝 Submissions")
    student_files = st.file_uploader("Upload Student Files", type=['py', 'zip', 'pdf', 'docx', 'txt'], accept_multiple_files=True, key="students")

# Reset confirmation ONLY if files actually change
current_files = {
    "rubric": rubric_file.name if rubric_file else None,
    "sol": sol_file.name if sol_file else None,
    "students": len(student_files) if student_files else 0
}

if current_files != st.session_state.file_states:
    st.session_state.confirm_no_sol = False
    st.session_state.show_confirm_dialog = False
    st.session_state.file_states = current_files
    st.session_state.run_evaluation = False

# ==========================================
# 5. GRADING ENGINE
# ==========================================
# Handle Run button click
col_btn_1, col_btn_2, col_btn_3 = st.columns([1, 2, 1])
with col_btn_2:
    if st.button("🚀 Run Evaluation", type="primary", use_container_width=True):
        # Reset run flags
        st.session_state.run_evaluation = True
        # If no solution and not already confirmed, we need to ask
        if not sol_file and not st.session_state.confirm_no_sol:
            st.session_state.show_confirm_dialog = True
            st.session_state.run_evaluation = False # Wait for confirm

# Show Confirmation Dialog if needed
if st.session_state.get("show_confirm_dialog", False):
    st.warning("⚠️ No Solution file uploaded. The AI will grade based only on the Rubric and its own judgment, which may be less accurate.")
    c1, c2 = st.columns(2)
    if c1.button("Confirm: Proceed without Solution"):
        st.session_state.confirm_no_sol = True
        st.session_state.show_confirm_dialog = False
        st.session_state.run_evaluation = True # Trigger run
        st.rerun()
    if c2.button("Cancel"):
        st.session_state.show_confirm_dialog = False
        st.rerun()

# Safety checks and execution
if st.session_state.run_evaluation:
    num_students = len(student_files)
    
    if balance < num_students:
        st.error(f"Insufficient credits! You need {num_students} but have {balance}.")
        st.session_state.run_evaluation = False
    elif not (rubric_file and student_files):
        st.error("Please upload Rubric and Student Files before running.")
        st.session_state.run_evaluation = False
    else:
        # Proceed with evaluation
        try:
            client = genai.Client(api_key=api_key)
            rubric_text = extract_text(rubric_file)
            
            if sol_file:
                sol_text = extract_text(sol_file)
                context_content = f"RUBRIC: {rubric_text}\n\nOFFICIAL SOLUTION: {sol_text}"
            else:
                context_content = f"RUBRIC: {rubric_text}\n\n(No official solution provided. Please use your best judgment as an expert teaching assistant to evaluate the student's work based strictly on the rubric provided.)"

            all_results = []
            raw_list = []
            zip_buffer = io.BytesIO()
            successfully_graded = 0
            
            # Setup JSON Schema
            response_schema = {
                "type": "OBJECT",
                "properties": {
                    "username": {"type": "STRING"},
                    "total_score": {"type": "NUMBER"},
                    "overall_feedback": {"type": "STRING"},
                    "questions": {
                        "type": "ARRAY",
                        "items": {
                            "type": "OBJECT",
                            "properties": {
                                "q_num": {"type": "NUMBER"},
                                "score": {"type": "NUMBER"},
                                "feedback": {"type": "STRING"}
                            },
                            "required": ["q_num", "score", "feedback"]
                        }
                    }
                },
                "required": ["username", "total_score", "overall_feedback", "questions"]
            }

            progress_bar = st.progress(0)
            status_text = st.empty()

            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                for i, f in enumerate(student_files):
                    status_text.info(f"Grading {i+1}/{num_students}: {f.name}")
                    content = extract_text(f)
                    
                    prompt = f"{context_content}\n\nStudent: {f.name}\nWork: {content}"
                    
                    try:
                        response = client.models.generate_content(
                            model=model_choice,
                            contents=prompt,
                            config=types.GenerateContentConfig(
                                response_mime_type='application/json',
                                response_schema=response_schema
                            )
                        )
                        data = json.loads(response.text)
                        raw_list.append(data)
                        
                        # Process for Gradebook
                        flat_entry = {"Username": data.get("username"), "Total Score": data.get("total_score")}
                        for q in data.get("questions", []):
                            flat_entry[f"Q{int(q.get('q_num'))} Score"] = q.get('score')
                        all_results.append(flat_entry)
                        
                        # Generate PDF for ZIP
                        pdf_bytes = create_pdf_report(data)
                        zip_file.writestr(f"{data['username']}_Report.pdf", pdf_bytes)
                        successfully_graded += 1
                        
                    except Exception as e:
                        st.error(f"Error grading {f.name}: {e}")
                    
                    progress_bar.progress((i + 1) / num_students)

            # SAVE RESULTS & DEDUCT CREDITS
            if successfully_graded > 0:
                deduct_credit(user_email, successfully_graded)
                df_results = pd.DataFrame(all_results)
                st.session_state.grading_data = df_results
                st.session_state.raw_json_list = raw_list
                st.session_state.zip_bytes = zip_buffer.getvalue()
                
                csv_buffer = io.StringIO()
                df_results.to_csv(csv_buffer, index=False)
                st.session_state.csv_data = csv_buffer.getvalue()
                
                st.rerun()

        except Exception as global_e:
            st.error(f"Critical System Error: {global_e}")
        finally:
            st.session_state.run_evaluation = False # Reset flag after run

# ==========================================
# 6. RESULTS DISPLAY (TABS)
# ==========================================
if st.session_state.grading_data is not None:
    st.divider()
    t1, t2, t3 = st.tabs(["📊 Summary", "📈 Analytics", "👤 Reports"])
    
    with t1:
        df = st.session_state.grading_data
        st.dataframe(df, use_container_width=True)
        c1, c2 = st.columns(2)
        c1.download_button("📂 Download PDFs (ZIP)", st.session_state.zip_bytes, "reports.zip", use_container_width=True)
        c2.download_button("📊 Download Gradebook (CSV)", st.session_state.csv_data, "grades.csv", use_container_width=True)

    with t2:
        st.subheader("Class Score Distribution")
        if "Total Score" in df.columns:
            st.bar_chart(df.set_index("Username")["Total Score"])

    with t3:
        # Individual report viewer with arrows
        u_list = [s["username"] for s in st.session_state.raw_json_list]
        
        col_prev, col_sel, col_next = st.columns([1, 4, 1])
        
        with col_prev:
            if st.button("⬅️ Previous", use_container_width=True):
                st.session_state.report_index = (st.session_state.report_index - 1) % len(u_list)
        
        with col_next:
            if st.button("Next ➡️", use_container_width=True):
                st.session_state.report_index = (st.session_state.report_index + 1) % len(u_list)
        
        with col_sel:
            # Update index if user selects from dropdown manually
            selected = st.selectbox("Select Student", u_list, index=st.session_state.report_index)
            st.session_state.report_index = u_list.index(selected)

        # Display feedback for the selected student
        s_data = st.session_state.raw_json_list[st.session_state.report_index]
        st.write(f"## Report for: **{s_data['username']}**")
        st.metric("Total Score", f"{s_data['total_score']} pts")
        
        st.subheader("Overall Feedback")
        st.info(s_data['overall_feedback'])
        
        st.subheader("Question-by-Question Breakdown")
        for q in s_data.get("questions", []):
            with st.expander(f"Q{q.get('q_num')} | Score: {q.get('score')}", expanded=True):
                st.write(q.get('feedback', 'No feedback provided.'))

# ==========================================
# 7. GETTING STARTED
# ==========================================
st.divider()
st.header("📖 Getting Started")
st.write("New to the AI Teaching Assistant? Watch this quick guide to learn how to upload your rubric, solution, and student submissions for automated grading.")
st.video("https://www.youtube.com/watch?v=dQw4w9WgXcQ") # Placeholder video link

